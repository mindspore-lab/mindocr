import re
from html.parser import HTMLParser

import docx
from bs4 import BeautifulSoup
from docx import Document


def get_table_rows(table_soup):
    """
    Get all rows for the table.
    """
    table_row_selectors = [
        "table > tr",
        "table > thead > tr",
        "table > tbody > tr",
        "table > tfoot > tr",
    ]
    return table_soup.select(", ".join(table_row_selectors), recursive=False)


def get_table_columns(row):
    """
    Get all columns for the specified row tag.
    """
    return row.find_all(["th", "td"], recursive=False) if row else []


def get_table_dimensions(table_soup):
    """
    Get the number of rows and columns in the table.
    """
    rows = get_table_rows(table_soup)
    cols = get_table_columns(rows[0]) if rows else []

    col_count = 0
    for col in cols:
        colspan = col.attrs.get("colspan", 1)
        col_count += int(colspan)

    return rows, col_count


def get_cell_html(soup):
    """
    Return the HTML content of a cell without the <td> tags.
    """
    return " ".join([str(i) for i in soup.contents])


def delete_paragraph(paragraph):
    """
    Delete a paragraph from a docx document.
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_whitespace(string, leading=False, trailing=False):
    """
    Remove white space from a string.
    """
    if leading:
        string = re.sub(r"^\s*\n+\s*", "", string)
    if trailing:
        string = re.sub(r"\s*\n+\s*$", "", string)
    string = re.sub(r"\s*\n\s*", " ", string)
    return re.sub(r"\s+", " ", string)


font_styles = {
    "b": "bold",
    "strong": "bold",
    "em": "italic",
    "i": "italic",
    "u": "underline",
    "s": "strike",
    "sup": "superscript",
    "sub": "subscript",
    "th": "bold",
}

font_names = {
    "code": "Courier",
    "pre": "Courier",
}


class HtmlToDocx(HTMLParser):
    def __init__(self):
        super().__init__()
        self.options = {
            "fix-html": True,
            "images": True,
            "tables": True,
            "styles": True,
        }
        self.table_row_selectors = [
            "table > tr",
            "table > thead > tr",
            "table > tbody > tr",
            "table > tfoot > tr",
        ]
        self.table_style = None
        self.paragraph_style = None

    def set_initial_attrs(self, document=None):
        self.tags = {
            "span": [],
            "list": [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options["fix-html"]
        self.document = self.doc
        self.include_tables = True
        self.include_images = self.options["images"]
        self.include_styles = self.options["styles"]
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        """
        Copy settings from another instance of HtmlToDocx
        """
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def ignore_nested_tables(self, tables_soup):
        """
        Return only the highest level tables.
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all("table"))
        return new_tables

    def get_tables(self):
        """
        Get all tables from the HTML.
        """
        if not hasattr(self, "soup"):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all("table"))
        self.table_no = 0

    def run_process(self, html):
        """
        Process the HTML content.
        """
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, "html.parser")
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_cell(self, html, cell):
        """
        Add HTML content to a table cell.
        """
        if not isinstance(cell, docx.table._Cell):
            raise ValueError("Second argument needs to be a %s" % docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        if not self.doc.paragraphs:
            self.doc.add_paragraph("")

    def apply_paragraph_style(self, style=None):
        """
        Apply style to the current paragraph.
        """
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def handle_table(self, html, doc):
        """
        Handle nested tables by parsing them manually.
        """
        table_soup = BeautifulSoup(html, "html.parser")
        rows, cols_len = get_table_dimensions(table_soup)
        table = doc.add_table(len(rows), cols_len)
        table.style = doc.styles["Table Grid"]

        num_rows = len(table.rows)
        num_cols = len(table.columns)

        cell_row = 0
        for _, row in enumerate(rows):
            cols = get_table_columns(row)
            cell_col = 0
            for col in cols:
                colspan = int(col.attrs.get("colspan", 1))
                rowspan = int(col.attrs.get("rowspan", 1))

                cell_html = get_cell_html(col)
                if col.name == "th":
                    cell_html = f"<b>{cell_html}</b>"

                if cell_row >= num_rows or cell_col >= num_cols:
                    continue

                docx_cell = table.cell(cell_row, cell_col)

                while docx_cell.text != "":  # Skip the merged cell
                    cell_col += 1
                    docx_cell = table.cell(cell_row, cell_col)

                cell_to_merge = table.cell(cell_row + rowspan - 1, cell_col + colspan - 1)
                if docx_cell != cell_to_merge:
                    docx_cell.merge(cell_to_merge)

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html or " ", docx_cell)

                cell_col += colspan
            cell_row += 1

    def handle_data(self, data):
        """
        Handle text data within HTML tags.
        """
        if self.skip:
            return

        if "pre" not in self.tags:
            data = remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        link = self.tags.get("a")
        if link:
            self.handle_link(link["href"], data)
        else:
            self.run = self.paragraph.add_run(data)
            spans = self.tags["span"]
            for span in spans:
                if "style" in span:
                    style = self.parse_dict_string(span["style"])
                    self.add_styles_to_run(style)

            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name
