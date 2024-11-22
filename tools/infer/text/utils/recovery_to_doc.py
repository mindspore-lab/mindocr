import os
from typing import Dict, List

from docx import Document, shared
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .table_process import HtmlToDocx


def set_document_styles(doc: Document) -> None:
    """
    Set the styles for the document.
    Args:
        doc (Document): The document to set styles for.
    """
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = shared.Pt(6.5)


def convert_info_docx(res: List[Dict], save_folder: str, doc_name: str) -> None:
    """
    Convert OCR results to a DOCX file.
    Args:
        res (List[Dict]): OCR results.
        save_folder (str): Folder to save the DOCX file.
        doc_name (str): Name of the DOCX file.
    Returns:
        None
    """
    doc = Document()
    set_document_styles(doc)

    flag = 1  # Current layout flag
    previous_layout = None  # To record the previous layout

    for region in res:
        if not region["res"]:
            continue

        # Check if the current layout has changed to avoid creating the same layout repeatedly
        current_layout = region["layout"]
        if current_layout != previous_layout:
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            if current_layout == "single":
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
                flag = 1
            elif current_layout == "double":
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
                flag = 2
            elif current_layout == "triple":
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "3")
                flag = 3
            previous_layout = current_layout  # Update the previous layout record

        # Insert content based on the region type
        if region["type"].lower() == "figure":
            img_path = region["res"]
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            # Insert picture, width depends on the column layout
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2.5))
            elif flag == 3:
                run.add_picture(img_path, width=shared.Inches(1.5))

        elif region["type"].lower() == "title":
            doc.add_heading(region["res"])

        elif region["type"].lower() == "table":
            parser = HtmlToDocx()
            parser.table_style = "TableGrid"
            parser.handle_table(region["res"], doc)

        else:  # Default to handling text regions
            paragraph = doc.add_paragraph()
            text_run = paragraph.add_run(region["res"])
            text_run.font.size = shared.Pt(10)

    # Save as DOCX file
    docx_path = os.path.join(save_folder, f"{doc_name}_ocr.docx")
    doc.save(docx_path)


def sorted_layout_boxes(res: List[Dict], w: int) -> List[Dict]:
    """
    Sort boxes based on distribution, supporting single, double, and triple column layouts,
    considering columns with large spans.
    Args:
        res (List[Dict]): Results from layout.
        w (int): Document width.
    Returns:
        List[Dict]: Sorted results.
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]["layout"] = "single"
        return res

    # Sort by y-coordinate from top to bottom, then by x-coordinate from right to left
    sorted_boxes = sorted(res, key=lambda x: (x["bbox"][1], -x["bbox"][0]))
    _boxes = list(sorted_boxes)

    res_left = []
    res_center = []
    res_right = []
    new_res = []

    column_thresholds = [w / 3, 2 * w / 3]
    tolerance = 0.02 * w

    # First round: classify columns, determine the distribution of boxes in each column
    for current_box in _boxes:
        box_left, box_right = current_box["bbox"][0], current_box["bbox"][2]
        box_width = box_right - box_left

        # Determine column layout, ensuring each box is assigned to only one column
        if box_width > column_thresholds[1]:
            current_box["layout"] = "spanning"
            new_res.append(current_box)
        elif box_right < column_thresholds[0] + tolerance:
            res_left.append(current_box)
        elif box_left > column_thresholds[1] - tolerance:
            res_right.append(current_box)
        elif column_thresholds[0] - tolerance <= box_left <= column_thresholds[1] + tolerance:
            res_center.append(current_box)
        else:
            res_left.append(current_box)

    # Second round: determine specific layout based on column distribution
    for box in res_left:
        if res_center and res_right:
            box["layout"] = "triple"
        elif res_right or res_center:
            box["layout"] = "double"
        else:
            box["layout"] = "single"
        new_res.append(box)

    for box in res_center:
        box["layout"] = "triple" if res_left and res_right else "double"
        new_res.append(box)

    for box in res_right:
        box["layout"] = "triple" if res_center else "double"
        new_res.append(box)

    return new_res
